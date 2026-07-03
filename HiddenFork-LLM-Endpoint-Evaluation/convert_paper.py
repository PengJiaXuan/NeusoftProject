"""
Convert Hidden_Fork_v4_paper.md → DOCX + PDF (广东东软学院 thesis template).

Template specs (from template1-13.png):
  A4, margins: top/bottom 2.54 cm, left 3.17 cm, right 2.50 cm
  Body text: 12 pt (小四) Times New Roman + SimSun, 1.5× line spacing,
             first-line indent 2 characters (~0.85 cm)
  Chapter headings: 18 pt bold, centered (SimHei / TNR)
  Section headings: 14 pt bold, left-aligned (SimHei / TNR)
  Page header (body only): "广东东软学院本科毕业设计（论文）" 10.5 pt + bottom border
  Page footer (body only): centered  -N-  page numbers starting at 1
  Structure: Cover → 原创性声明 → 版权授权 → 摘要 → Abstract → 目录 →
             Body chapters → 参考文献 → 致谢 → 附录
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from fpdf import FPDF

ROOT = Path(__file__).parent
SRC = ROOT / "Hidden_Fork_v4_paper.md"
OUT_DOCX = ROOT / "Hidden_Fork_v4_thesis_v22.docx"
OUT_PDF  = ROOT / "Hidden_Fork_v4_thesis_v22.pdf"
md_text = SRC.read_text(encoding="utf-8")

# ═══════════════════════ Constants ═══════════════════════
SCHOOL = "广东东软学院"
TITLE_CN = "隐匿分叉：大语言模型在API、Web应用\n与影子中继端点间的行为差异测量"
TITLE_EN = ("Hidden Fork: Measuring Behavioral Divergence Across\n"
            "API, Web App, and Shadow Relay Endpoints\n"
            "for Large Language Models")

STUDENT_NAME = "彭嘉铉"
STUDENT_ID   = "22265120128"
DEPARTMENT   = "国际教育学院"
MAJOR        = "软件工程（中外合作办学）"
CLASS_NAME   = "22软件联合1班"
SUPERVISOR   = "任萌"
SUPERVISOR_TITLE = "助教"
COMPLETION_DATE  = "2026.04.15"

ABSTRACT_CN = (
    "大语言模型（LLM）评估通常假设通过一个接口访问的模型在通过另一个接口访问时表现一致。"
    "本文通过系统比较三个前沿大语言模型——GPT-5.4、Claude Sonnet 4.6和Gemini 3 Flash——"
    "在三种不同端点类型（官方API、官方Web应用和第三方影子中继）上的测量行为差异，"
    "对这一假设进行了实证检验。"
    "\n\n"
    "使用包含AIME 2025（n=30）、GPQA Diamond（n=100）、MedQA（n=100）、"
    "LegalBench（n=100）、MMLU-Pro（n=100）以及基于JailbreakBench的安全性评估"
    "（n=100）的基准测试套件，我们发现端点选择可以在能力和安全性指标上产生实质性差异。"
    "GPT-5.4在API和Web应用端点之间的AIME准确率下降了26.6个百分点。"
    "Claude Sonnet 4.6在不同基准和端点间表现出不一致的差异方向性。"
    "Gemini 3 Flash影子中继在AIME上产生了完全的评分失败（0.0%），"
    "归因于协议级交互而非模型能力缺失。安全拒绝率也因端点而异，"
    "其中Claude Sonnet 4.6在Web应用和影子端点上表现出比API更高的有害性评分。"
    "\n\n"
    "这些发现表明，端点选择不是一个表面的部署细节，而是影响大语言模型测量行为的"
    "实质性因素，对可重复性和基准解释具有重要意义。"
)
KEYWORDS_CN = "大语言模型；API评估；端点差异；影子中继；基准可重复性；测量表面"
KEYWORDS_EN = ("large language models; API evaluation; endpoint divergence; "
               "shadow relay; benchmark reproducibility; measurement surface")

DECLARATION_TEXT = (
    "本人郑重声明：所呈交的毕业设计（论文），是本人在指导老师的指导下，"
    "独立进行的设计（研究）工作及取得的成果。文中引用他人的文献、数据、图件、"
    "资料均已明确标注出，论文中的结论和结果为本人独立完成，不包含他人成果及为获得"
    "其他教育机构的学位或证书而使用其材料。与我一同工作的同志对本设计（研究）所做的"
    "任何贡献均已在论文中作了明确的说明并表示了致意。"
    "\n\n"
    "本声明的法律结果由本人承担。"
)

AUTHORIZATION_TEXT = (
    "本学位论文作者完全了解广东东软学院有关保留、使用学位论文的规定。"
    "本人同意学校有权保留并向国家有关部门或机构送交论文的复印件和电子版，"
    "允许论文被查阅和借阅。本人授权广东东软学院可以将学位论文的全部或部分内容"
    "编入有关数据库进行检索，可以采用影印、缩印或扫描等复制手段保存和汇编本学位论文。"
)

AI_STATEMENT_TEXT = (
    "本人郑重声明：本毕业设计（论文）研究及撰写全过程，核心研究思路、实验设计、"
    "逻辑架构、数据分析与观点结论，均由本人在指导教师指导下独立完成。"
    "论文创作过程中，仅在文献资料整理、语句语法润色、格式规范校对、"
    "专业术语释义等辅助环节有限使用生成式人工智能工具；"
    "论文主体内容、研究创新点、核心论证与成果表述，均未由人工智能直接生成、代写或替代创作。"
    "\n\n"
    "凡论文中引用的文献、数据、图表及参考观点均已规范标注出处，"
    "未利用生成式人工智能工具进行抄袭、拼凑学术内容、伪造研究数据与成果等学术不端行为。"
    "若存在不实声明及违规使用情形，本人自愿承担相应学术责任与一切相关后果。"
)

ACK_CN = (
    "在整个毕业设计的各个阶段我都得到了任萌老师的悉心指导和耐心帮助。"
    "任老师在实验设计与学术写作方面的建设性意见对本论文的最终成形起到了关键作用，"
    "在此向任老师表示最衷心的感谢。"
    "\n\n"
    "同时感谢国际教育学院的同学们在研究过程中给予的有益讨论和支持。"
    "特别感谢本研究所使用的开源工具、基准数据集和评估框架的开发者和维护者，"
    "他们的公开工作使本研究得以顺利进行。"
    "\n\n"
    "最后，我要感谢家人在我本科学习期间给予的始终如一的支持和理解。"
)

ACK_EN = (
    "I would like to express my sincere gratitude to my supervisor, Ms. Ren Meng, "
    "for her continuous guidance, patience, and encouragement throughout the course "
    "of this research. Her constructive feedback on experimental design and academic "
    "writing was instrumental in shaping this thesis into its final form."
    "\n\n"
    "I am also grateful to my classmates and friends in the School of International "
    "Education who provided helpful discussions and moral support during the research "
    "process. Special thanks go to the developers and maintainers of the open-source "
    "tools, benchmark datasets, and evaluation frameworks used in this study, whose "
    "publicly available work made this research possible."
    "\n\n"
    "Finally, I wish to thank my family for their unwavering support and understanding "
    "throughout my undergraduate studies at Neusoft Institute Guangdong."
)


# ═══════════════════════ Markdown Parser ═══════════════════════
def parse_md(text):
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]; stripped = line.strip()

        # Headings
        if stripped.startswith("### "):
            yield ("h3", stripped[4:]); i += 1; continue
        if stripped.startswith("## "):
            yield ("h2", stripped[3:]); i += 1; continue
        if stripped.startswith("# "):
            yield ("h1", stripped[2:]); i += 1; continue

        # HR
        if stripped == "---":
            i += 1; continue

        # Single-line equation  $...$
        if stripped.startswith("$") and stripped.endswith("$") and len(stripped) > 4 and "$$" not in stripped:
            yield ("equation", stripped); i += 1; continue

        # Display equation  $$...$$
        if stripped.startswith("$$"):
            if stripped.endswith("$$") and len(stripped) > 4:
                yield ("equation", stripped); i += 1; continue
            buf = [stripped[2:]]; i += 1
            while i < len(lines):
                if lines[i].strip().endswith("$$"):
                    buf.append(lines[i].strip()[:-2]); i += 1; break
                buf.append(lines[i].strip()); i += 1
            yield ("equation", " ".join(buf)); continue

        # Image
        m = re.match(r"^!\[(.+?)\]\((.+?)\)", stripped)
        if m:
            yield ("image", (m.group(1), m.group(2))); i += 1; continue

        # Table
        if "|" in stripped and stripped.startswith("|"):
            header = [c.strip() for c in stripped.split("|") if c.strip()]
            i += 1
            if i < len(lines) and re.match(r"^\|[\s\-:|]+\|", lines[i].strip()):
                i += 1  # skip separator
            rows = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                rows.append([c.strip() for c in lines[i].split("|") if c.strip()])
                i += 1
            yield ("table", (header, rows)); continue

        # Blank
        if not stripped:
            i += 1; continue

        # Paragraph — collect continuation lines
        para = [stripped]; i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (not nxt or nxt.startswith("#") or nxt.startswith("|") or
                nxt == "---" or nxt.startswith("$$") or re.match(r"^!\[", nxt) or
                (nxt.startswith("$") and nxt.endswith("$") and len(nxt) > 4)):
                break
            para.append(nxt); i += 1
        yield ("para", " ".join(para))


def clean_md(text):
    t = re.sub(r"\*\*\*(.+?)\*\*\*", r"\1", text)
    t = re.sub(r"\*\*(.+?)\*\*", r"\1", t)
    t = re.sub(r"\*(.+?)\*", r"\1", t)
    t = re.sub(r"`(.+?)`", r"\1", t)
    return t


def resolve_img(p):
    full = ROOT / p
    return full if full.exists() else None


def latex_to_readable(raw):
    t = raw.strip().strip("$").strip()
    t = re.sub(r"\\text\{([^}]*)\}", r"\1", t)
    t = re.sub(r"\\frac\{([^}]*)\}\{([^}]*)\}", r"(\1) / (\2)", t)
    t = re.sub(r"\\max_\{([^}]*)\}", r"max over \1 of ", t)
    t = re.sub(r"\\mathcal\{([^}]*)\}", r"\1", t)
    t = re.sub(r"\|([^|]+)\|", r"|\1|", t)
    t = re.sub(r"_\{([^}]*)\}", r"_\1", t)
    t = re.sub(r"\s+", " ", t).strip()
    return t


# ═══════════════════════ Parse all content ═══════════════════════
all_blocks = list(parse_md(md_text))

abstract_en = ""
body_blocks = []
ref_blocks = []
in_body = False
in_refs = False
skip_paper_title = True

for kind, data in all_blocks:
    if kind == "h1" and skip_paper_title:
        skip_paper_title = False; continue
    if kind == "h2" and data.strip().lower() == "abstract":
        in_body = False; continue
    if kind == "h2" and re.match(r"^\d+\.\s", data):
        in_body = True; in_refs = False
    if kind == "h2" and "References" in data:
        in_body = False; in_refs = True; continue
    if in_refs:
        ref_blocks.append((kind, data)); continue
    if not in_body and kind == "para" and "i noticed something odd" in data.lower() and not abstract_en:
        abstract_en = clean_md(data); continue
    if in_body:
        body_blocks.append((kind, data))


# ═══════════════════════ DOCX Helpers ═══════════════════════
print("Generating DOCX ...")
doc = Document()

# --- Page setup (default section = cover + preliminary pages) ---
for sec in doc.sections:
    sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17); sec.right_margin = Cm(2.50)
    sec.header_distance = Cm(1.50); sec.footer_distance = Cm(1.75)
    sec.header_distance = Cm(1.50)
    sec.footer_distance = Cm(1.75)

# --- Base styles ---
sn = doc.styles["Normal"]
sn.font.name = "Times New Roman"; sn.font.size = Pt(12)
sn.paragraph_format.line_spacing = 1.5
sn.paragraph_format.space_after = Pt(0); sn.paragraph_format.space_before = Pt(0)
# Set East Asian font on Normal style
rpr = sn.element.rPr
if rpr is None:
    rpr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>'); sn.element.append(rpr)
ea = rpr.find(qn("w:rFonts"))
if ea is None:
    ea = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="SimSun"/>'); rpr.append(ea)
else:
    ea.set(qn("w:eastAsia"), "SimSun")

# Heading styles
for lv, sz, align in [(1, 16, WD_ALIGN_PARAGRAPH.CENTER),
                       (2, 15, WD_ALIGN_PARAGRAPH.LEFT),
                       (3, 14, WD_ALIGN_PARAGRAPH.LEFT)]:
    hs = doc.styles[f"Heading {lv}"]
    hs.font.name = "Arial"; hs.font.size = Pt(sz)
    hs.font.color.rgb = RGBColor(0, 0, 0); hs.font.bold = True
    hs.paragraph_format.line_spacing = Pt(20)
    hs.paragraph_format.alignment = align
    if lv == 1:
        hs.paragraph_format.space_before = Pt(6)
        hs.paragraph_format.space_after = Pt(6)
    elif lv == 2:
        hs.paragraph_format.space_before = Pt(6)
        hs.paragraph_format.space_after = Pt(6)
    # Set east-asia font to SimHei for headings
    hs_rpr = hs.element.rPr
    if hs_rpr is None:
        hs_rpr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>'); hs.element.append(hs_rpr)
    hs_ea = hs_rpr.find(qn("w:rFonts"))
    if hs_ea is None:
        hs_ea = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="SimHei"/>'); hs_rpr.append(hs_ea)
    else:
        hs_ea.set(qn("w:eastAsia"), "SimHei")


def set_run_font(run, latin="Times New Roman", east_asia="SimSun",
                 size=Pt(12), bold=False, italic=False, underline=False):
    run.font.name = latin; run.font.size = size
    run.bold = bold; run.italic = italic; run.underline = underline
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), east_asia)
    rFonts.set(qn("w:ascii"), latin)
    rFonts.set(qn("w:hAnsi"), latin)


def set_char_indent(para, first_line_chars=None, left_chars=None, hanging_chars=None):
    """Set paragraph indent in *character* units (so 论无忧 recognizes it)."""
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    # Clear conflicting absolute values
    for attr in ("firstLine", "hanging", "left", "start"):
        if ind.get(qn(f"w:{attr}")) is not None:
            ind.attrib.pop(qn(f"w:{attr}"))
    if first_line_chars is not None:
        ind.set(qn("w:firstLineChars"), str(first_line_chars))
    if left_chars is not None:
        ind.set(qn("w:leftChars"), str(left_chars))
    if hanging_chars is not None:
        ind.set(qn("w:hangingChars"), str(hanging_chars))


def set_para_spacing(para, before_pt=None, after_pt=None):
    """Set space-before / space-after in points."""
    pPr = para._p.get_or_add_pPr()
    sp = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        pPr.append(sp)
    if before_pt is not None:
        sp.set(qn("w:before"), str(int(before_pt * 20)))
        sp.set(qn("w:beforeLines"), "0")
    if after_pt is not None:
        sp.set(qn("w:after"), str(int(after_pt * 20)))
        sp.set(qn("w:afterLines"), "0")


def set_char_indent(p, first_line_chars=None, left_chars=None, hanging_chars=None):
    """Add character-based indent XML attributes so format checkers recognise them."""
    pPr = p._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    if first_line_chars is not None:
        ind.set(qn("w:firstLineChars"), str(first_line_chars))
    if left_chars is not None:
        ind.set(qn("w:leftChars"), str(left_chars))
    if hanging_chars is not None:
        ind.set(qn("w:hangingChars"), str(hanging_chars))
        # hanging = negative firstLine
        if first_line_chars is None:
            ind.set(qn("w:firstLineChars"), str(-hanging_chars))


def dp(text, bold=False, italic=False, size=None, align=None,
       indent=True, east_asia="SimSun"):
    """Add a body paragraph."""
    p = doc.add_paragraph()
    if indent:
        set_char_indent(p, first_line_chars=200)
    p.paragraph_format.line_spacing = Pt(20)
    p.alignment = align if align else WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    set_run_font(r, east_asia=east_asia, size=Pt(size or 12), bold=bold, italic=italic)
    return p


def dp_cn(text, bold=False, size=None, indent=True):
    """Add a Chinese paragraph (SimSun body, SimHei bold)."""
    ea = "SimHei" if bold else "SimSun"
    return dp(text, bold=bold, size=size, indent=indent, east_asia=ea)


def blank(n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0


def centered(text, size=12, bold=False, east_asia="SimSun"):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    set_run_font(r, east_asia=east_asia, size=Pt(size), bold=bold)
    return p


def title_line(text, size=18, bold=True, east_asia="SimHei"):
    """Justified title line with first-line indent (论无忧 requirement)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.0
    set_char_indent(p, first_line_chars=200)
    r = p.add_run(text)
    set_run_font(r, east_asia=east_asia, size=Pt(size), bold=bold)
    return p


def centered_multi(text, size=12, bold=False, east_asia="SimSun"):
    """Centered text; split on \\n into separate paragraphs (no soft breaks)."""
    for i, line in enumerate(text.split("\n")):
        if line.strip():
            centered(line.strip(), size, bold, east_asia)


def add_heading_cn(text, level=1):
    """Add a heading with Chinese east-asia font (SimHei). No blank lines added — caller manages spacing."""
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(h, before_pt=6, after_pt=6)
    if level == 2:
        set_char_indent(h, left_chars=200)
    for run in h.runs:
        set_run_font(run, latin="Arial", east_asia="SimHei",
                     size=Pt(16 if level == 1 else 15), bold=True)
    return h


def add_docx_table(header, rows):
    nc = len(header)
    t = doc.add_table(rows=1 + len(rows), cols=nc)
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, h in enumerate(header):
        c = t.rows[0].cells[ci]; c.text = ""
        pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.paragraph_format.line_spacing = Pt(20)
        rr = pp.add_run(h); rr.bold = True
        set_run_font(rr, size=Pt(10.5), bold=True)
        sh = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F0F0" w:val="clear"/>')
        c._tc.get_or_add_tcPr().append(sh)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            if ci >= nc: break
            c = t.rows[1 + ri].cells[ci]; c.text = ""
            pp = c.paragraphs[0]
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            pp.paragraph_format.line_spacing = Pt(20)
            rr = pp.add_run(val)
            set_run_font(rr, size=Pt(10.5))
    doc.add_paragraph()


def add_docx_image(alt, path):
    res = resolve_img(path)
    if res and res.suffix.lower() == ".png":
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.add_run().add_picture(str(res), width=Inches(5.3))
        c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_after = Pt(6)
        # Normalize "Figure X.X. Desc" -> "Figure X.X Desc" and strip trailing dot
        cap = re.sub(r'^((?:Figure|Table)\s+\d+(?:\.\d+)?)\.(?=\s)', r'\1', alt)
        rr = c.add_run(cap)
        set_run_font(rr, size=Pt(10.5), italic=True)


eq_num = 0

def add_docx_equation(raw):
    global eq_num
    eq_num += 1
    readable = latex_to_readable(raw)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(8)
    # Center tab + right tab so equation centers and number right-aligns
    from docx.shared import Cm as _Cm
    from docx.enum.text import WD_TAB_ALIGNMENT
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(_Cm(7.5), WD_TAB_ALIGNMENT.CENTER)
    tab_stops.add_tab_stop(_Cm(15.3), WD_TAB_ALIGNMENT.RIGHT)
    r0 = p.add_run("\t")
    r = p.add_run(readable)
    set_run_font(r, size=Pt(12), italic=True)
    r1 = p.add_run("\t")
    r2 = p.add_run(f"({eq_num})")
    set_run_font(r2, size=Pt(12))


def add_page_field(para, fmt="-"):
    """Insert  -PAGE-  field into a paragraph."""
    r1 = para.add_run(fmt)
    set_run_font(r1, size=Pt(10.5))
    # PAGE field begin
    run_b = para.add_run()
    fc_begin = OxmlElement("w:fldChar")
    fc_begin.set(qn("w:fldCharType"), "begin")
    run_b._r.append(fc_begin)
    # instrText
    run_i = para.add_run()
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = " PAGE "
    run_i._r.append(it)
    # separate
    run_s = para.add_run()
    fc_sep = OxmlElement("w:fldChar")
    fc_sep.set(qn("w:fldCharType"), "separate")
    run_s._r.append(fc_sep)
    # placeholder
    run_n = para.add_run("1")
    set_run_font(run_n, size=Pt(10.5))
    # end
    run_e = para.add_run()
    fc_end = OxmlElement("w:fldChar")
    fc_end.set(qn("w:fldCharType"), "end")
    run_e._r.append(fc_end)
    # trailing dash
    r2 = para.add_run(fmt)
    set_run_font(r2, size=Pt(10.5))


def add_header_border(para):
    """Add a bottom border line to a header paragraph."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ═══════════════════════ Cover Page ═══════════════════════
blank(2)
centered(SCHOOL, 22, True, "SimHei")
blank(1)
centered("Neusoft Institute Guangdong", 18, True)
blank(1)
centered("本科毕业设计（论文）", 22, True, "SimHei")
centered("Undergraduate Dissertation", 18, True)
blank(2)
for line in TITLE_CN.split("\n"):
    if line.strip():
        title_line(line.strip(), 18, True, "SimHei")
blank(1)
for line in TITLE_EN.split("\n"):
    if line.strip():
        title_line(line.strip(), 18, True, "SimSun")
blank(3)

# Cover fields as a 2-column table
cover_fields = [
    ("学    院", DEPARTMENT),
    ("专    业", MAJOR),
    ("班    级", CLASS_NAME),
    ("学    号", STUDENT_ID),
    ("学生姓名", STUDENT_NAME),
    ("指导教师", SUPERVISOR),
    ("导师职称", SUPERVISOR_TITLE),
    ("完成日期", COMPLETION_DATE),
]
cover_table = doc.add_table(rows=len(cover_fields), cols=2)
cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cover_table.autofit = False
cover_table.style = "Table Grid"

# Add borders to table (Table Grid style should do this, but enforce explicitly)
tbl = cover_table._tbl
tblPr = tbl.tblPr
tblBorders = OxmlElement("w:tblBorders")
for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
    b = OxmlElement(f"w:{border_name}")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), "6")
    b.set(qn("w:color"), "000000")
    tblBorders.append(b)
tblPr.append(tblBorders)

for i, (label, value) in enumerate(cover_fields):
    row = cover_table.rows[i]
    row.height = Cm(1.0)
    # Label cell — LEFT aligned label keeps the form readable
    c1 = row.cells[0]
    c1.width = Cm(4.0)
    c1.text = ""
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p1.paragraph_format.line_spacing = 1.0
    # 班级 row: 段前 0.5 行
    if label.strip() == "班    级":
        set_para_spacing(p1, before_pt=6)
    r1 = p1.add_run(label)
    set_run_font(r1, east_asia="SimSun", size=Pt(15), bold=True)
    # Value cell — CENTER aligned per 论无忧 requirement
    c2 = row.cells[1]
    c2.width = Cm(9.0)
    c2.text = ""
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.line_spacing = 1.0
    if label.strip() == "班    级":
        set_para_spacing(p2, before_pt=6)
    r2 = p2.add_run(value)
    set_run_font(r2, east_asia="SimSun", size=Pt(15))

doc.add_page_break()

# ═══════════════════════ Declaration (原创性声明) ═══════════════════════
blank(1)
centered("毕业设计（论文）原创性声明", 18, True, "SimHei")
blank(1)
for para_text in DECLARATION_TEXT.split("\n\n"):
    dp_cn(para_text.strip())
blank(3)
p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
r = p.add_run("毕业设计文作者（签字）：                    签字日期：2026.04.15")
set_run_font(r, east_asia="SimSun", size=Pt(12))
doc.add_page_break()

# ═══════════════════════ Authorization (版权使用授权声明) ═══════════════════════
blank(1)
centered("毕业设计（论文）版权使用授权声明", 18, True, "SimHei")
blank(1)
for para_text in AUTHORIZATION_TEXT.split("\n\n"):
    dp_cn(para_text.strip())
blank(3)
p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
r = p.add_run("本学位论文作者签名：              指导老师签名：")
set_run_font(r, east_asia="SimSun", size=Pt(12))
blank(1)
p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
r = p.add_run("日期：2026.04.15")
set_run_font(r, east_asia="SimSun", size=Pt(12))
doc.add_page_break()

# ═══════════════════════ AI Use Statement (生成式人工智能工具使用声明) ═══════════════════════
blank(1)
centered("生成式人工智能工具使用声明", 18, True, "SimHei")
blank(1)
for para_text in AI_STATEMENT_TEXT.split("\n\n"):
    dp_cn(para_text.strip())
blank(3)
p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
r = p.add_run("毕业论文作者签名：")
set_run_font(r, east_asia="SimSun", size=Pt(12))
blank(1)
p = doc.add_paragraph()
p.paragraph_format.line_spacing = 1.5
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("日　期：2026.04.15")
set_run_font(r, east_asia="SimSun", size=Pt(12))
doc.add_page_break()

# ═══════════════════════ Chinese Abstract (摘要) ═══════════════════════
blank(1)
p_abs_title = centered("摘　　要", 16, True, "SimHei")
p_abs_title.paragraph_format.line_spacing = 1.0
blank(1)
for para_text in ABSTRACT_CN.split("\n\n"):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.99)
    set_char_indent(p, first_line_chars=200)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(para_text.strip())
    set_run_font(r, east_asia="SimSun", size=Pt(14))
blank(1)
p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.5
set_char_indent(p, left_chars=286, hanging_chars=286)
r = p.add_run("关键词：")
set_run_font(r, east_asia="SimHei", size=Pt(14), bold=True)
r2 = p.add_run(KEYWORDS_CN)
set_run_font(r2, east_asia="SimSun", size=Pt(14))
doc.add_page_break()

# ═══════════════════════ English Abstract ═══════════════════════
blank(1)
p_abs_en_title = centered("Abstract", 16, True)
p_abs_en_title.paragraph_format.line_spacing = 1.5
blank(1)
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(0.99)
set_char_indent(p, first_line_chars=200)
p.paragraph_format.line_spacing = 1.5
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r = p.add_run(abstract_en)
set_run_font(r, size=Pt(14))
blank(1)
p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.5
set_char_indent(p, left_chars=286, hanging_chars=286)
r = p.add_run("Key words: ")
set_run_font(r, size=Pt(14), bold=True)
r2 = p.add_run(KEYWORDS_EN)
set_run_font(r2, size=Pt(14))
doc.add_page_break()

# ═══════════════════════ TOC — English (Table of Contents) ═══════════════════════
blank(1)
p_toc_title = doc.add_paragraph()
p_toc_title.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
set_char_indent(p_toc_title, first_line_chars=200)
r_toc = p_toc_title.add_run("Contents")
set_run_font(r_toc, latin="Arial", east_asia="SimHei", size=Pt(16), bold=True)
blank(1)

# TOC field code
def add_toc_field():
    p_toc = doc.add_paragraph()
    run_toc = p_toc.add_run()
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run_toc._r.append(fld_char_begin)
    run_instr = p_toc.add_run()
    instr_text = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve">'
        ' TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    run_instr._r.append(instr_text)
    run_sep = p_toc.add_run()
    fld_char_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    run_sep._r.append(fld_char_sep)
    run_placeholder = p_toc.add_run(
        "(Open in Word, press Ctrl+A then F9 to generate TOC)")
    set_run_font(run_placeholder, east_asia="SimSun", size=Pt(10.5))
    run_placeholder.font.color.rgb = RGBColor(128, 128, 128)
    run_end = p_toc.add_run()
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_end._r.append(fld_char_end)

add_toc_field()

# updateFields setting
settings = doc.settings.element
update = OxmlElement("w:updateFields")
update.set(qn("w:val"), "true")
settings.append(update)


# ═══════════════════════ Section break → Body ═══════════════════════
# Create a new section for body pages with header + footer
from docx.enum.section import WD_SECTION_START
sec_body = doc.add_section(WD_SECTION_START.NEW_PAGE)
sec_body.page_width = Cm(21.0); sec_body.page_height = Cm(29.7)
sec_body.top_margin = Cm(2.54); sec_body.bottom_margin = Cm(2.54)
sec_body.left_margin = Cm(3.17); sec_body.right_margin = Cm(2.50)
sec_body.header_distance = Cm(1.50); sec_body.footer_distance = Cm(1.75)

# --- Header ---
sec_body.header.is_linked_to_previous = False
hdr_para = sec_body.header.paragraphs[0]
hdr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_run = hdr_para.add_run("Neusoft Institute Guangdong Undergraduate Dissertation")
set_run_font(hdr_run, latin="Times New Roman", east_asia="SimSun", size=Pt(9))
add_header_border(hdr_para)

# --- Footer with -N- page numbers ---
sec_body.footer.is_linked_to_previous = False
ftr_para = sec_body.footer.paragraphs[0]
ftr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_page_field(ftr_para, "-")

# --- Reset page number to 1 ---
sectPr = sec_body._sectPr
pgNumType = OxmlElement("w:pgNumType")
pgNumType.set(qn("w:start"), "1")
sectPr.append(pgNumType)


# ═══════════════════════ Body Chapters ═══════════════════════
chapter_num = 0
section_num = 0
first_chapter = True
just_after_chapter = False

for kind, data in body_blocks:
    if kind == "h2":
        chapter_num += 1; section_num = 0
        title = re.sub(r"^\d+\.\s*", "", data).strip()
        if first_chapter:
            first_chapter = False
        else:
            doc.add_page_break()
        blank(1)
        h = add_heading_cn(f"Chapter {chapter_num} {title}", level=1)
        blank(1)
        just_after_chapter = True

    elif kind == "h3":
        section_num += 1
        title = re.sub(r"^\d+\.\d+\s*", "", data).strip()
        title = title.rstrip("?")
        if not just_after_chapter:
            blank(1)
        just_after_chapter = False
        h = doc.add_heading(f"{chapter_num}.{section_num} {title}", level=2)
        set_para_spacing(h, before_pt=6, after_pt=6)
        set_char_indent(h, left_chars=200)
        for run in h.runs:
            set_run_font(run, latin="Arial", east_asia="SimHei", size=Pt(15), bold=True)

    elif kind == "image":
        just_after_chapter = False
        add_docx_image(data[0], data[1])

    elif kind == "table":
        just_after_chapter = False
        add_docx_table(data[0], data[1])

    elif kind == "equation":
        just_after_chapter = False
        add_docx_equation(data)

    elif kind == "para":
        just_after_chapter = False
        cleaned = clean_md(data)
        # Table/figure caption lines should be centered, no period at end
        if re.match(r'^(Table|Figure)\s+\d+', cleaned):
            # Remove period ONLY if it's after the number AND followed by whitespace.
            # Use lookahead to avoid backtracking that would eat the "4.1" dot.
            cleaned = re.sub(r'^((?:Table|Figure)\s+\d+(?:\.\d+)?)\.(?=\s)', r'\1', cleaned)
            cap_size = 10.5 if cleaned.startswith("Figure") else 12
            dp(cleaned, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=cap_size)
        else:
            dp(cleaned)


# ═══════════════════════ References (参考文献) ═══════════════════════
doc.add_page_break()
blank(1)
h = add_heading_cn("References", level=1)
# Override to 0pt as 论无忧 expects for References heading
set_para_spacing(h, before_pt=0, after_pt=0)
blank(1)
for kind, data in ref_blocks:
    if kind == "para":
        p = doc.add_paragraph(); p.paragraph_format.line_spacing = Pt(20)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        ref_text = clean_md(data).rstrip()
        if not ref_text.endswith("."):
            ref_text += "."
        r = p.add_run(ref_text)
        set_run_font(r, size=Pt(10.5))


# ═══════════════════════ Appendix (附录) ═══════════════════════
doc.add_page_break()
h = add_heading_cn("Appendix", level=1)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = Pt(20)
r = p.add_run("A. Quality Gate Summary")
set_run_font(r, latin="Times New Roman", east_asia="SimHei", size=Pt(12), bold=True)
dp("The final merged dataset comprises 54 benchmark-endpoint cells "
   "(3 models × 3 endpoints × 6 benchmarks). All 54 cells passed "
   "quality-gate validation with zero errors and zero missing raw "
   "responses (54/54 ok).")

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = Pt(20)
r = p.add_run("B. Run Selection Manifest")
set_run_font(r, latin="Times New Roman", east_asia="SimHei", size=Pt(12), bold=True)
dp("The following table summarizes the source run for each model-endpoint "
   "slice in the final merged dataset.")
manifest = [
    ["GPT-5.4", "API", "formal_main_v1"],
    ["GPT-5.4", "App", "formal_main_v1"],
    ["GPT-5.4", "Shadow", "formal_shadow_rerun_v1"],
    ["Claude Sonnet 4.6", "API", "formal_main_v1"],
    ["Claude Sonnet 4.6", "App", "formal_claude_app_max20_v1"],
    ["Claude Sonnet 4.6", "Shadow", "formal_shadow_rerun_v1"],
    ["Gemini 3 Flash", "API", "formal_gemini_api_vertex_v1"],
    ["Gemini 3 Flash", "App", "formal_main_v1"],
    ["Gemini 3 Flash", "Shadow", "formal_shadow_rerun_v1"],
]
add_docx_table(["Model", "Endpoint", "Source Run"], manifest)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.line_spacing = Pt(20)
r = p.add_run("C. Gemini Shadow AIME Failure Audit Detail")
set_run_font(r, latin="Times New Roman", east_asia="SimHei", size=Pt(12), bold=True)
dp("A per-item audit of all 30 AIME 2025 responses from the Gemini 3 Flash "
   "shadow relay is available in the project repository. The audit confirms "
   "that 30/30 responses were abruptly truncated, 0/30 contained a valid "
   "final-answer marker, and the scorer extracted incidental tail integers "
   "from unfinished reasoning in 27/30 cases.")


# ═══════════════════════ Acknowledgements (致谢) ═══════════════════════
doc.add_page_break()
h = add_heading_cn("Acknowledgements", level=1)
for para_text in ACK_CN.split("\n\n"):
    dp_cn(para_text.strip())
blank(1)
for para_text in ACK_EN.split("\n\n"):
    dp(para_text.strip())


doc.save(str(OUT_DOCX))
print(f"  -> {OUT_DOCX}")


# ═══════════════════════ PDF ═══════════════════════
print("Generating PDF ...")

class PaperPDF(FPDF):
    def __init__(self, *a, **kw):
        super().__init__(*a, **kw)
        self.in_body = False

    def header(self):
        if not self.in_body or self.page <= 1:
            return
        self.set_font("TNR", "", 9)
        self.cell(0, 6, "Neusoft Institute Guangdong Undergraduate Dissertation", align="C")
        self.set_draw_color(0, 0, 0)
        y = self.get_y() + 7
        self.line(self.l_margin, y, self.w - self.r_margin, y)
        self.ln(10)

    def footer(self):
        if not self.in_body:
            return
        self.set_y(-15)
        self.set_font("TNR", "", 10.5)
        self.cell(0, 10, f"-{self.page_no()}-", align="C")


pdf = PaperPDF("P", "mm", "A4")
pdf.set_auto_page_break(True, 25)

# Fonts
FD = Path("C:/Windows/Fonts")
for s, f in [("", "times.ttf"), ("B", "timesbd.ttf"),
             ("I", "timesi.ttf"), ("BI", "timesbi.ttf")]:
    pdf.add_font("TNR", s, str(FD / f))

# Chinese fonts
try:
    pdf.add_font("SimSun", "", str(FD / "simsun.ttc"))
    pdf.add_font("SimHei", "", str(FD / "simhei.ttf"))
    has_cn_font = True
except Exception:
    has_cn_font = False
    print("  [warn] Chinese fonts not loaded; PDF will use TNR for Chinese text")

LM, RM, TM = 31.7, 25.0, 25.4
pdf.set_margins(LM, TM, RM)
LH = 6.5; BS = 12


def aw():
    return pdf.w - pdf.l_margin - pdf.r_margin


def pcn(text, size=12, bold=False, align="J", indent=True):
    """Print Chinese paragraph."""
    fname = "SimHei" if (bold and has_cn_font) else ("SimSun" if has_cn_font else "TNR")
    style = "B" if (bold and not has_cn_font) else ""
    pdf.set_font(fname, style, size)
    if indent:
        pdf.set_x(pdf.l_margin + 7); w = aw() - 7
    else:
        w = aw()
    pdf.multi_cell(w, LH, text, align=align)
    pdf.ln(1.5)


def pen(text, size=12, bold=False, italic=False, align="J", indent=True):
    """Print English paragraph."""
    style = ("B" if bold else "") + ("I" if italic else "")
    pdf.set_font("TNR", style, size)
    if indent:
        pdf.set_x(pdf.l_margin + 7); w = aw() - 7
    else:
        w = aw()
    pdf.multi_cell(w, LH, clean_md(text), align=align)
    pdf.ln(1.5)


def ph1(text, cn=False):
    pdf.add_page()
    fname = "SimHei" if (cn and has_cn_font) else "TNR"
    style = "B" if not (cn and has_cn_font) else ""
    pdf.set_font(fname, style, 18)
    pdf.multi_cell(0, 10, text, align="C")
    pdf.ln(6)


def ph2(text):
    pdf.ln(4); pdf.set_font("TNR", "B", 14)
    pdf.multi_cell(0, 8, text); pdf.ln(3)


def ptbl(hdr, rows):
    nc = len(hdr); cw = aw() / nc
    need = (1 + len(rows)) * 5 + 10
    if pdf.get_y() + need > pdf.h - 30:
        pdf.add_page()
    pdf.set_font("TNR", "B", 9)
    for h in hdr:
        pdf.cell(cw, 5, h, border=1, align="C")
    pdf.ln()
    pdf.set_font("TNR", "", 9)
    for row in rows:
        for ci, v in enumerate(row):
            pdf.cell(cw, 5, v if ci < len(row) else "", border=1,
                     align="L" if ci == 0 else "C")
        pdf.ln()
    pdf.ln(3)


def pimg(alt, path):
    res = resolve_img(path)
    if not res or not res.exists():
        pen(f"[Image: {alt}]", italic=True, size=9); return
    if pdf.get_y() + 100 > pdf.h - 30:
        pdf.add_page()
    pdf.image(str(res), x=pdf.l_margin, w=aw())
    pdf.ln(3)
    pdf.set_font("TNR", "I", 9)
    pdf.multi_cell(0, 4, alt, align="C")
    pdf.ln(3)


# ────── PDF Cover ──────
pdf.add_page()
pdf.ln(20)
if has_cn_font:
    pdf.set_font("SimHei", "", 22)
    pdf.cell(0, 12, SCHOOL, align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(6)
    pdf.set_font("SimHei", "", 22)
    pdf.cell(0, 12, "本科毕业设计（论文）", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(15)
    pdf.set_font("SimHei", "", 16)
    pdf.multi_cell(0, 9, TITLE_CN.replace("\n", ""), align="C")
else:
    pdf.set_font("TNR", "B", 22)
    pdf.cell(0, 12, "Neusoft Institute Guangdong", align="C",
             new_x="LMARGIN", new_y="NEXT")
    pdf.ln(6)
    pdf.set_font("TNR", "B", 18)
    pdf.cell(0, 12, "Undergraduate Graduation Thesis", align="C",
             new_x="LMARGIN", new_y="NEXT")
    pdf.ln(15)
pdf.ln(4)
pdf.set_font("TNR", "B", 14)
pdf.multi_cell(0, 8, TITLE_EN.replace("\n", " "), align="C")
pdf.ln(20)
if has_cn_font:
    pdf.set_font("SimSun", "", 14)
    for lab, val in [("系 / 院：", DEPARTMENT), ("专业名称：", MAJOR),
                     ("学　　号：", STUDENT_ID), ("学生姓名：", STUDENT_NAME),
                     ("学习形式：", "全日制"), ("指导教师：", SUPERVISOR),
                     ("完成时间：", "2026 年 3 月")]:
        pdf.cell(0, 8, f"{lab}  {val}", align="C",
                 new_x="LMARGIN", new_y="NEXT")
else:
    pdf.set_font("TNR", "", 13)
    for lab, val in [("Student Name:", STUDENT_NAME), ("Student ID:", STUDENT_ID),
                     ("Department:", DEPARTMENT), ("Major:", MAJOR),
                     ("Supervisor:", SUPERVISOR), ("Date:", "March 2026")]:
        pdf.cell(0, 8, f"{lab}  {val}", align="C",
                 new_x="LMARGIN", new_y="NEXT")

# ────── PDF Declaration ──────
if has_cn_font:
    ph1("毕业设计（论文）原创性声明", cn=True)
    for pt in DECLARATION_TEXT.split("\n\n"):
        pcn(pt.strip())
    pdf.ln(15)
    pcn("毕业设计文作者（签字）：           签字日期：   年  月  日", indent=False)
else:
    ph1("Declaration of Originality")
    pen("This thesis is entirely my own work.")

# ────── PDF Authorization ──────
if has_cn_font:
    ph1("毕业设计（论文）版权使用授权声明", cn=True)
    for pt in AUTHORIZATION_TEXT.split("\n\n"):
        pcn(pt.strip())
    pdf.ln(15)
    pcn("本学位论文作者签名：          指导老师签名：", indent=False)
else:
    ph1("Authorization for Use")
    pen("Authorization text.")

# ────── PDF Chinese Abstract ──────
if has_cn_font:
    ph1("摘　　要", cn=True)
    for pt in ABSTRACT_CN.split("\n\n"):
        pcn(pt.strip())
    pdf.ln(4)
    pcn(f"关键词：{KEYWORDS_CN}", bold=True, indent=False)

# ────── PDF English Abstract ──────
ph1("Abstract")
pen(abstract_en)
pdf.ln(4)
pdf.set_font("TNR", "B", BS)
pdf.cell(22, LH, "Keywords: ")
pdf.set_font("TNR", "", BS)
pdf.multi_cell(0, LH, KEYWORDS_EN)

# ────── PDF Body ──────
pdf.in_body = True
cn_num = 0; sn_ = 0; pdf_eq_num = 0

for kind, data in body_blocks:
    if kind == "h2":
        cn_num += 1; sn_ = 0
        title = re.sub(r"^\d+\.\s*", "", data).strip()
        ph1(f"Chapter {cn_num}  {title}")
    elif kind == "h3":
        sn_ += 1
        title = re.sub(r"^\d+\.\d+\s*", "", data).strip()
        ph2(f"{cn_num}.{sn_}  {title}")
    elif kind == "image":
        pimg(data[0], data[1])
    elif kind == "table":
        ptbl(data[0], data[1])
    elif kind == "equation":
        pdf_eq_num += 1
        pdf.ln(3); pdf.set_font("TNR", "I", 12)
        readable = latex_to_readable(data)
        pdf.cell(0, LH, f"{readable}          ({pdf_eq_num})",
                 align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(3)
    elif kind == "para":
        pen(data)

# ────── PDF References ──────
if has_cn_font:
    ph1("参考文献", cn=True)
else:
    ph1("References")
for kind, data in ref_blocks:
    if kind == "para":
        pen(clean_md(data), indent=False, size=10.5)

# ────── PDF Acknowledgements ──────
if has_cn_font:
    ph1("致　　谢", cn=True)
    for pt in ACK_CN.split("\n\n"):
        pcn(pt.strip())
    pdf.ln(4)
else:
    ph1("Acknowledgements")
for pt in ACK_EN.split("\n\n"):
    pen(pt.strip())

# ────── PDF Appendix ──────
if has_cn_font:
    ph1("附　　录", cn=True)
else:
    ph1("Appendix")
pen("A. Quality Gate Summary", bold=True, indent=False)
pen("The final merged dataset comprises 54 benchmark-endpoint cells. "
    "All 54 cells passed quality-gate validation (54/54 ok).")
pen("B. Run Selection Manifest", bold=True, indent=False)
pen("See DOCX version for full manifest table.")
pen("C. Gemini Shadow AIME Failure Audit Detail", bold=True, indent=False)
pen("30/30 responses truncated, 0/30 valid final answers, "
    "27/30 incidental tail integer extractions (all incorrect).")

pdf.output(str(OUT_PDF))
print(f"  -> {OUT_PDF}")
print("Done.")
